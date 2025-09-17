"""
Redact PII in DOCX or TXT using Presidio.

Usage (DOCX mode, defaults to the repo's assets file):
  python redact_docx.py
  python redact_docx.py --input "assets/Qualia Case Study.docx" --output "redaction/Qualia Case Study_redacted.docx"

Usage (TXT mode - single file):
  python redact_docx.py --txt-file "path/to/file.txt"
  # Output is created in ./redaction as file_redacted.txt

Usage (TXT mode - directory of .txt files):
  python redact_docx.py --txt-dir "path/to/folder"
  # Each *.txt is processed to ./redaction/*_redacted.txt

Tag options (apply to all modes):
  python redact_docx.py --org-tag "<organisation>"
  python redact_docx.py --person-tag "<person>" --email-tag "<email>" --phone-tag "<phone>"

Simple behavior: PERSON -> <person>, ORG/ORGANIZATION -> <organisation>, EMAIL_ADDRESS -> <email>, PHONE_NUMBER -> <phone>, GPE/LOC/LOCATION -> <LOC>, NORP -> <COUNTRY>, countries are preserved as-is, everything else -> <redacted>.
"""

from __future__ import annotations

from pathlib import Path
import argparse
import re
import re

from docx import Document  # python-docx
from presidio_analyzer import AnalyzerEngine, RecognizerResult

# Normalized lowercase country names and common aliases
COUNTRY_ALIASES = {
    "afghanistan","albania","algeria","andorra","angola","antigua and barbuda","argentina","armenia",
    "australia","austria","azerbaijan","bahamas","bahrain","bangladesh","barbados","belarus","belgium",
    "belize","benin","bhutan","bolivia","bosnia and herzegovina","botswana","brazil","brunei","bulgaria",
    "burkina faso","burundi","cabo verde","cape verde","cambodia","cameroon","canada","central african republic",
    "chad","chile","china","colombia","comoros","congo","democratic republic of the congo","republic of the congo",
    "costa rica","cote d'ivoire","côte d'ivoire","croatia","cuba","cyprus","czech republic","czechia",
    "denmark","djibouti","dominica","dominican republic","ecuador","egypt","el salvador","equatorial guinea",
    "eritrea","estonia","eswatini","swaziland","ethiopia","fiji","finland","france","gabon","gambia",
    "georgia","germany","ghana","greece","grenada","guatemala","guinea","guinea-bissau","guyana","haiti",
    "honduras","hungary","iceland","india","indonesia","iran","iraq","ireland","israel","italy","jamaica",
    "japan","jordan","kazakhstan","kenya","kiribati","kosovo","kuwait","kyrgyzstan","laos","lao pdr","latvia",
    "lebanon","lesotho","liberia","libya","liechtenstein","lithuania","luxembourg","madagascar","malawi",
    "malaysia","maldives","mali","malta","marshall islands","mauritania","mauritius","mexico","micronesia",
    "moldova","monaco","mongolia","montenegro","morocco","mozambique","myanmar","burma","namibia","nauru",
    "nepal","netherlands","new zealand","nicaragua","niger","nigeria","north korea","democratic people's republic of korea",
    "north macedonia","macedonia","norway","oman","pakistan","palau","panama","papua new guinea","paraguay","peru",
    "philippines","poland","portugal","qatar","romania","russia","russian federation","rwanda","saint kitts and nevis",
    "saint lucia","saint vincent and the grenadines","samoa","san marino","sao tome and principe","são tomé and príncipe",
    "saudi arabia","senegal","serbia","seychelles","sierra leone","singapore","slovakia","slovenia","solomon islands",
    "somalia","south africa","south korea","republic of korea","korea, republic of","spain","sri lanka","sudan",
    "south sudan","suriname","sweden","switzerland","syria","syrian arab republic","taiwan","chinese taipei",
    "tajikistan","tanzania","thailand","timor-leste","east timor","togo","tonga","trinidad and tobago","tunisia",
    "turkey","türkiye","turkmenistan","tuvalu","uganda","ukraine","united arab emirates","uae","united kingdom",
    "great britain","britain","uk","u.k","u.k.","united states","united states of america","usa","u.s.a","u.s.a.",
    "us","u.s","u.s.","uruguay","uzbekistan","vanuatu","venezuela","vietnam","viet nam","yemen","zambia","zimbabwe",
    "hong kong","macau","macao","palestine","state of palestine","vatican city","holy see","western sahara",
}
from presidio_anonymizer import AnonymizerEngine, OperatorConfig


def redact_text(
    text: str,
    analyzer: AnalyzerEngine,
    anonymizer: AnonymizerEngine,
    operators: dict,
) -> str:
    """Analyze text for PII and return anonymized text using provided operators."""
    if not text or not text.strip():
        return text

    try:
        results = analyzer.analyze(text=text, language="en")
    except Exception as exc:  # Provide a clear hint if spaCy model is missing
        raise RuntimeError(
            "Presidio Analyzer failed. Ensure spaCy model 'en_core_web_lg' is installed: "
            "python -m spacy download en_core_web_lg"
        ) from exc

    if not results:
        return text

    # Retag country mentions among geo entities so they use COUNTRY label (we'll preserve them)
    def _retag_countries(results_list):
        retagged: list[RecognizerResult] = []
        for r in results_list:
            if r.entity_type in {"GPE", "LOC", "LOCATION"}:
                span_text = text[r.start : r.end].strip().lower().strip(".,;:()[]{}\"'")
                # Look for contextual cues like 'country' near the span
                before = text[max(0, r.start - 24) : r.start].lower()
                after = text[r.end : min(len(text), r.end + 24)].lower()
                has_country_cue = (
                    ("country" in before) or ("country" in after) or ("nation" in before) or ("nation" in after)
                )
                # Also check a comprehensive alias set (EN/FR/ES variants)
                if has_country_cue or span_text in COUNTRY_ALIASES:
                    retagged.append(
                        RecognizerResult(
                            entity_type="COUNTRY", start=r.start, end=r.end, score=r.score
                        )
                    )
                    continue
            retagged.append(r)
        return retagged

    results = _retag_countries(results)

    # Do not redact standalone years like 2010, 1999, 2024
    year_re = re.compile(r"^(?:19|20)\d{2}$")
    filtered_results: list[RecognizerResult] = []
    for r in results:
        if r.entity_type == "DATE_TIME":
            span = text[r.start : r.end].strip().strip(".,;:()[]{}\"'")
            if year_re.match(span):
                continue
        filtered_results.append(r)
    results = filtered_results

    # Preserve countries as-is by removing them from anonymization
    results = [r for r in results if r.entity_type != "COUNTRY"]

    anonymized = anonymizer.anonymize(
        text=text,
        analyzer_results=results,
        operators=operators,
    )
    return anonymized.text


_BOLD_PREFIX_RE = re.compile(r'^(\*\*.+?\*\*:\s*)(.*)$')


def redact_text_preserving_bold_prefix(
    text: str,
    analyzer: AnalyzerEngine,
    anonymizer: AnonymizerEngine,
    operators: dict,
) -> str:
    """Redact text per line, preserving bold prefixes like '**Title**: ' at start of line.

    - If a line starts with '**...**: ', that prefix is kept verbatim and only the remainder is redacted.
    - Other lines are fully redacted.
    """
    if not text:
        return text

    # Fast path for single-line strings
    if '\n' not in text and '\r' not in text:
        m = _BOLD_PREFIX_RE.match(text)
        if m:
            prefix, rest = m.group(1), m.group(2)
            return prefix + redact_text(rest, analyzer, anonymizer, operators)
        return redact_text(text, analyzer, anonymizer, operators)

    # Multi-line processing preserving exact newline characters
    out_parts: list[str] = []
    for line in text.splitlines(keepends=True):
        # Separate newline from content
        newline = ''
        core = line
        if line.endswith('\r\n'):
            core = line[:-2]
            newline = '\r\n'
        elif line.endswith('\n') or line.endswith('\r'):
            core = line[:-1]
            newline = line[-1]

        m = _BOLD_PREFIX_RE.match(core)
        if m:
            prefix, rest = m.group(1), m.group(2)
            out_parts.append(prefix + redact_text(rest, analyzer, anonymizer, operators) + newline)
        else:
            out_parts.append(redact_text(core, analyzer, anonymizer, operators) + newline)

    return ''.join(out_parts)


def _build_presidio(org_tag: str, person_tag: str, email_tag: str, phone_tag: str):
    """Create analyzer/anonymizer and configured operators (shared by modes)."""
    analyzer = AnalyzerEngine()
    anonymizer = AnonymizerEngine()
    operators = {
        "DEFAULT": OperatorConfig("replace", {"new_value": "<redacted>"}),
        "PERSON": OperatorConfig("replace", {"new_value": person_tag}),
        "ORG": OperatorConfig("replace", {"new_value": org_tag}),
        "ORGANIZATION": OperatorConfig("replace", {"new_value": org_tag}),
        # Geo-like entities
        "GPE": OperatorConfig("replace", {"new_value": "<LOC>"}),
        "LOC": OperatorConfig("replace", {"new_value": "<LOC>"}),
        "LOCATION": OperatorConfig("replace", {"new_value": "<LOC>"}),
        "COUNTRY": OperatorConfig("replace", {"new_value": "<COUNTRY>"}),
        "NORP": OperatorConfig("replace", {"new_value": "<COUNTRY>"}),
        "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": email_tag}),
        "PHONE_NUMBER": OperatorConfig("replace", {"new_value": phone_tag}),
    }
    return analyzer, anonymizer, operators


def redact_docx(
    input_path: Path,
    output_path: Path,
    org_tag: str,
    person_tag: str,
    email_tag: str,
    phone_tag: str,
) -> None:
    """Load DOCX, redact paragraphs and table cells, and save to output."""
    analyzer, anonymizer, operators = _build_presidio(org_tag, person_tag, email_tag, phone_tag)

    doc = Document(str(input_path))

    # Redact standalone paragraphs
    for paragraph in doc.paragraphs:
        paragraph.text = redact_text_preserving_bold_prefix(paragraph.text, analyzer, anonymizer, operators)

    # Redact text inside tables (common in reports)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = redact_text_preserving_bold_prefix(paragraph.text, analyzer, anonymizer, operators)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def redact_text_file(
    input_path: Path,
    output_path: Path,
    org_tag: str,
    person_tag: str,
    email_tag: str,
    phone_tag: str,
) -> None:
    """Load a TXT file, redact its text, and write to output."""
    analyzer, anonymizer, operators = _build_presidio(org_tag, person_tag, email_tag, phone_tag)
    text = input_path.read_text(encoding="utf-8")
    redacted = redact_text_preserving_bold_prefix(text, analyzer, anonymizer, operators)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(redacted, encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Redact PII from DOCX or TXT using Presidio.")
    parser.add_argument(
        "--input",
        type=Path,
        default=Path("assets") / "Qualia Case Study.docx",
        help="Input DOCX path (default: assets/Qualia Case Study.docx)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Optional output DOCX path. Defaults to ./redaction/<input_stem>_redacted.docx",
    )
    parser.add_argument(
        "--txt-file",
        type=Path,
        default=None,
        help="Process a single TXT file. Output is <name>_redacted.txt beside input.",
    )
    parser.add_argument(
        "--txt-dir",
        type=Path,
        default=None,
        help="Process all *.txt files in this directory. Outputs are *_redacted.txt beside inputs.",
    )
    parser.add_argument(
        "--org-tag",
        type=str,
        default="<organisation>",
        help="Replacement tag for organisation names (default: <organisation>)",
    )
    parser.add_argument(
        "--person-tag",
        type=str,
        default="<person>",
        help="Replacement tag for person names (default: <person>)",
    )
    parser.add_argument(
        "--email-tag",
        type=str,
        default="<email>",
        help="Replacement tag for email addresses (default: <email>)",
    )
    parser.add_argument(
        "--phone-tag",
        type=str,
        default="<phone>",
        help="Replacement tag for phone numbers (default: <phone>)",
    )
    args = parser.parse_args()

    # TXT directory mode
    if args.txt_dir is not None:
        if not args.txt_dir.exists() or not args.txt_dir.is_dir():
            raise FileNotFoundError(f"TXT directory not found: {args.txt_dir}")
        out_root = Path("redaction")
        out_root.mkdir(parents=True, exist_ok=True)
        for txt_path in sorted(args.txt_dir.glob("*.txt")):
            out_path = out_root / f"{txt_path.stem}_redacted.txt"
            redact_text_file(
                txt_path,
                out_path,
                args.org_tag,
                args.person_tag,
                args.email_tag,
                args.phone_tag,
            )
            print(f"Redacted text saved to: {out_path}")
        return

    # TXT single-file mode
    if args.txt_file is not None:
        if not args.txt_file.exists():
            raise FileNotFoundError(f"TXT file not found: {args.txt_file}")
        out_root = Path("redaction")
        out_root.mkdir(parents=True, exist_ok=True)
        out_path = out_root / f"{args.txt_file.stem}_redacted.txt"
        redact_text_file(
            args.txt_file,
            out_path,
            args.org_tag,
            args.person_tag,
            args.email_tag,
            args.phone_tag,
        )
        print(f"Redacted text saved to: {out_path}")
        return

    # Default DOCX mode
    if not args.input.exists():
        raise FileNotFoundError(f"Input file not found: {args.input}")

    # Compute default DOCX output path in ./redaction
    out_docx = args.output
    if out_docx is None:
        Path("redaction").mkdir(parents=True, exist_ok=True)
        out_docx = Path("redaction") / f"{args.input.stem}_redacted.docx"

    redact_docx(
        args.input,
        out_docx,
        args.org_tag,
        args.person_tag,
        args.email_tag,
        args.phone_tag,
    )
    print(f"Redacted document saved to: {out_docx}")


if __name__ == "__main__":
    main()


